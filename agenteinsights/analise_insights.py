"""
Agente Insights - Análises e Chat IA
====================================
Versão: 1.5.0
Release: 5
Data: 02/06/2025

Descrição:
Executa análises (descritiva, preditiva, diagnóstica) e fornece função para chat IA,
onde o usuário pode consultar insights sobre tribos/squads, com respostas baseadas nos dados,
gráficos e tabelas. Permite salvar o histórico do chat em DOCX.
"""

import pandas as pd
import numpy as np
import os
import sys
import logging
import traceback
import re
import time
import json
import unicodedata
from datetime import datetime
from typing import Dict, Any, List, Tuple, Optional
from pathlib import Path
from collections import Counter, defaultdict
from sklearn.linear_model import LinearRegression
from sklearn.cluster import KMeans
from sklearn.preprocessing import StandardScaler
from docx import Document
from dotenv import load_dotenv
from openai import OpenAI
from unidecode import unidecode
import matplotlib.pyplot as plt
import seaborn as sns
import difflib
import torch
import cupy as cp  # Para operações GPU
from concurrent.futures import ThreadPoolExecutor

# Definir caminhos
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))  # Sobe um nível
DATA_DIR = os.path.join(BASE_DIR, 'dados')  # Usa 'dados' em vez de 'data'
OUTPUT_DIR = os.path.join(BASE_DIR, 'output')
GRAFICOS_DIR = os.path.join(OUTPUT_DIR, 'graficos')
RELATORIOS_DIR = os.path.join(OUTPUT_DIR, 'relatorios')

# Criar diretórios necessários
os.makedirs(DATA_DIR, exist_ok=True)
os.makedirs(GRAFICOS_DIR, exist_ok=True)
os.makedirs(RELATORIOS_DIR, exist_ok=True)

# Caminhos para arquivos
ARQUIVO_MATURIDADE = os.path.join(DATA_DIR, 'MaturidadeT.xlsx')
ARQUIVO_ALOCACAO = os.path.join(DATA_DIR, 'Alocacao.xlsx')
ARQUIVO_EXECUTIVO = os.path.join(DATA_DIR, 'Executivo.xlsx')

# Logging dos caminhos no início
logging.info(f"Arquivo de maturidade: {ARQUIVO_MATURIDADE}")
logging.info(f"Arquivo de alocação: {ARQUIVO_ALOCACAO}")
logging.info(f"Arquivo executivo: {ARQUIVO_EXECUTIVO}")

def normalizar_coluna(col):
    # Remove acentos, espaços e deixa minúsculo
    col = unicodedata.normalize('NFKD', str(col)).encode('ASCII', 'ignore').decode('ASCII')
    return col.strip().lower().replace(' ', '')

def mapear_colunas(df, nomes_esperados):
    # Cria um dicionário de {nome_normalizado: nome_original}
    col_map = {normalizar_coluna(c): c for c in df.columns}
    resultado = {}
    for nome in nomes_esperados:
        norm = normalizar_coluna(nome)
        if norm in col_map:
            resultado[nome] = col_map[norm]
        else:
            resultado[nome] = None
    return resultado

def carregar_dados() -> Dict[str, pd.DataFrame]:
    """Carrega os dados dos arquivos Excel"""
    dados = {}
    try:
        # Carregar arquivo de maturidade
        logging.info(f"Tentando carregar arquivo: {ARQUIVO_MATURIDADE}")
        if os.path.exists(ARQUIVO_MATURIDADE):
            df_maturidade = pd.read_excel(ARQUIVO_MATURIDADE)
            dados['maturidade'] = df_maturidade
            logging.info(f"Carregado arquivo MaturidadeT.xlsx com sucesso: {len(df_maturidade)} linhas")
            print(f"MaturidadeT.xlsx: {len(df_maturidade)} linhas")
            print(df_maturidade.head())
        else:
            logging.error(f"Arquivo não encontrado: {ARQUIVO_MATURIDADE}")
            return {}
        # Carregar arquivo de alocação
        logging.info(f"Tentando carregar arquivo: {ARQUIVO_ALOCACAO}")
        if os.path.exists(ARQUIVO_ALOCACAO):
            df_alocacao = pd.read_excel(ARQUIVO_ALOCACAO)
            dados['alocacao'] = df_alocacao
            logging.info(f"Carregado arquivo Alocacao.xlsx com sucesso: {len(df_alocacao)} linhas")
            print(f"Alocacao.xlsx: {len(df_alocacao)} linhas")
            print(df_alocacao.head())
        else:
            logging.error(f"Arquivo não encontrado: {ARQUIVO_ALOCACAO}")
            return {}
        # Carregar arquivo executivo (aba correta)
        logging.info(f"Tentando carregar arquivo: {ARQUIVO_EXECUTIVO}")
        if os.path.exists(ARQUIVO_EXECUTIVO):
            df_executivo = pd.read_excel(ARQUIVO_EXECUTIVO, sheet_name='NewBusinessAgility')
            dados['executivo'] = df_executivo
            logging.info(f"Carregado arquivo Executivo.xlsx com sucesso: {len(df_executivo)} linhas")
            print(f"Executivo.xlsx: {len(df_executivo)} linhas")
            print(df_executivo.head())
        else:
            logging.error(f"Arquivo não encontrado: {ARQUIVO_EXECUTIVO}")
            return {}
        if len(dados) == 3:
            logging.info("Todos os arquivos carregados com sucesso")
            return dados
        else:
            logging.error("Não foi possível carregar todos os arquivos necessários")
            return {}
    except Exception as e:
        logging.error(f"Erro ao carregar dados: {str(e)}")
        traceback.print_exc()
        return {}

def padronizar_ids(dados: Dict[str, pd.DataFrame]) -> Dict[str, pd.DataFrame]:
    """Padroniza os IDs e nomes das tribos para permitir o merge dos dados"""
    try:
        logging.info("Padronizando nomes das tribos para merge...")
        
        def limpar_nome_tribo(nome):
            if pd.isna(nome):
                return ''
            # Remove acentos e converte para minúsculas
            nome = unidecode(str(nome)).lower().strip()
            # Remove caracteres especiais
            nome = re.sub(r'[^a-z0-9\s]', '', nome)
            # Substitui múltiplos espaços por um único
            nome = re.sub(r'\s+', ' ', nome)
            return nome
            
        # Aplica a limpeza nos nomes das tribos
        for key, df in dados.items():
            if 'Tribo' in df.columns:
                df['nome_tribo_clean'] = df['Tribo'].apply(limpar_nome_tribo)
            elif 'tribe' in df.columns:
                df['nome_tribo_clean'] = df['tribe'].apply(limpar_nome_tribo)
                
        # Log dos nomes únicos em cada DataFrame
        nomes_maturidade = set(dados['maturidade']['nome_tribo_clean'].unique())
        nomes_alocacao = set(dados['alocacao']['nome_tribo_clean'].unique())
        
        logging.info(f"Nomes únicos em Maturidade: {sorted(list(nomes_maturidade))}")
        logging.info(f"Nomes únicos em Alocação: {sorted(list(nomes_alocacao))}")
        
        # Identifica nomes em comum
        nomes_comuns = nomes_maturidade.intersection(nomes_alocacao)
        logging.info(f"Total de nomes em comum: {len(nomes_comuns)}")
        logging.info(f"Nomes em comum: {nomes_comuns}")
        
        logging.info("IDs padronizados com sucesso")
        return dados
        
    except Exception as e:
        logging.error(f"Erro ao padronizar IDs: {str(e)}")
        traceback.print_exc()
        return {}

def cruzar_dados(dados: Dict[str, pd.DataFrame]) -> Dict[str, pd.DataFrame]:
    """Cruza os dados dos diferentes DataFrames"""
    try:
        logging.info("Iniciando cruzamento de dados...")
        
        # Extrai os DataFrames do dicionário
        df_maturidade = dados.get('maturidade')
        df_alocacao = dados.get('alocacao')
        
        if df_maturidade is None or df_alocacao is None:
            logging.error("DataFrames necessários não encontrados")
            return {}
            
        # Log das colunas disponíveis
        logging.info("Padronizando nomes das tribos para merge...")
        
        # Garante que temos a coluna nome_tribo_clean em ambos os DataFrames
        if 'nome_tribo_clean' not in df_maturidade.columns or 'nome_tribo_clean' not in df_alocacao.columns:
            logging.error("Coluna nome_tribo_clean ausente")
            return {}
            
        # Log dos nomes únicos em cada DataFrame
        nomes_maturidade = set(df_maturidade['nome_tribo_clean'].unique())
        nomes_alocacao = set(df_alocacao['nome_tribo_clean'].unique())
        
        logging.info(f"Nomes únicos em Maturidade: {sorted(list(nomes_maturidade))}")
        logging.info(f"Nomes únicos em Alocação: {sorted(list(nomes_alocacao))}")
        
        # Identifica nomes em comum
        nomes_comuns = nomes_maturidade.intersection(nomes_alocacao)
        logging.info(f"Total de nomes em comum: {len(nomes_comuns)}")
        logging.info(f"Nomes em comum: {nomes_comuns}")
        
        # Log das colunas antes do merge
        logging.info(f"Colunas em maturidade: {df_maturidade.columns.tolist()}")
        logging.info(f"Colunas em alocacao: {df_alocacao.columns.tolist()}")
        
        # Realiza o merge dos DataFrames
        df_consolidado = pd.merge(
            df_maturidade,
            df_alocacao,
            on='nome_tribo_clean',
            how='inner'
        )
        
        logging.info(f"Merge realizado com sucesso. Shape: {df_consolidado.shape}")
        
        # Log dos primeiros registros após o merge
        logging.info("\nPrimeiros registros após merge:")
        logging.info(f"\n{df_consolidado[['nome_tribo_clean', 'Tribo', 'Maturidade']].head()}")
        
        # Retorna um dicionário com os DataFrames originais e o consolidado
        return {
            'maturidade': df_maturidade,
            'alocacao': df_alocacao,
            'consolidado': df_consolidado
        }
        
    except Exception as e:
        logging.error(f"Falha ao cruzar dados: {str(e)}")
        traceback.print_exc()
        return {}

def executar_analises(df: pd.DataFrame) -> Dict[str, Any]:
    resultados = {}
    # Análise descritiva
    resultados['estatisticas'] = df.describe(include='all').to_dict()
    # Análise preditiva (regressão)
    num_cols = df.select_dtypes(include=[np.number]).columns
    if len(num_cols) > 1:
        X = df[num_cols[1:]].fillna(0)
        y = df[num_cols[0]].fillna(0)
        reg = LinearRegression().fit(X, y)
        resultados['regressao'] = {
            'coef': reg.coef_.tolist(),
            'intercept': float(reg.intercept_),
            'r2': reg.score(X, y)
        }
    else:
        resultados['regressao'] = None
    # Análise diagnóstica (clustering)
    if len(num_cols) > 1:
        X = df[num_cols].fillna(0)
        kmeans = KMeans(n_clusters=3, random_state=42).fit(X)
        resultados['clustering'] = {
            'labels': kmeans.labels_.tolist(),
            'centroids': kmeans.cluster_centers_.tolist(),
            'inertia': float(kmeans.inertia_)
        }
    else:
        resultados['clustering'] = None
    resultados['status'] = 'success'
    return resultados

def gerar_graficos(df: pd.DataFrame, resultados: Dict[str, Any]):
    """Gera gráficos para visualização dos dados"""
    try:
        # Criar diretório de gráficos se não existir
        os.makedirs(GRAFICOS_DIR, exist_ok=True)
        
        # Histograma da primeira coluna numérica
        colunas_numericas = df.select_dtypes(include=[np.number]).columns
        if len(colunas_numericas) > 0:
            plt.figure(figsize=(10, 6))
            sns.histplot(data=df, x=colunas_numericas[0])
            plt.title(f'Distribuição de {colunas_numericas[0]}')
            plt.savefig(os.path.join(GRAFICOS_DIR, 'histograma.png'))
        plt.close()
            
        # Box plot se houver dados suficientes
        if len(colunas_numericas) > 1:
            plt.figure(figsize=(12, 6))
            sns.boxplot(data=df[colunas_numericas[:3]])
            plt.title('Box Plot das Métricas Principais')
            plt.xticks(rotation=45)
            plt.tight_layout()
            plt.savefig(os.path.join(GRAFICOS_DIR, 'boxplot.png'))
            plt.close()
            
        # Gráfico de barras para métricas ágeis
        if 'metricas_ageis' in resultados:
            metricas = resultados['metricas_ageis']
            if isinstance(metricas, dict):
                plt.figure(figsize=(12, 6))
                valores = []
                labels = []
                for key, value in metricas.items():
                    if isinstance(value, (int, float)):
                        valores.append(value)
                        labels.append(key)
                if valores:
                    plt.bar(labels, valores)
                    plt.title('Métricas Ágeis')
                    plt.xticks(rotation=45)
                    plt.tight_layout()
                    plt.savefig(os.path.join(GRAFICOS_DIR, 'metricas_ageis.png'))
                    plt.close()
                    
    except Exception as e:
        logging.error(f"Erro ao gerar gráficos: {str(e)}")
        traceback.print_exc()

# Importações das funções do sistema
# As funções já estão definidas neste arquivo, não é necessário importá-las
# Correção Manual - Função chat_ia_loop
# ===================================
# Esta função corrigida deve substituir a função atual no arquivo analise_insights.py

"""
Implementação da função chat_ia_loop para interação com o usuário
"""

def chat_ia_loop(analises: List[Dict]):
    """
    Loop principal de interação com o usuário, mantendo contexto para perguntas de follow-up.
    """
    # Carrega variáveis de ambiente
    load_dotenv()
    
    # Inicializa cliente OpenAI
    client = openai.OpenAI()
    
    # Contexto base como uma mensagem de desenvolvedor
    contexto_base = {
        "role": "developer",
        "content": """
        # Identidade e Especialidades
        Você é um especialista em:
        * Gestão empresarial e organizacional
        * Agilidade e Business Agility
        * Gestão de projetos e produtos
        * Planejamento estratégico e indicadores
        
        # Instruções
        * Analise dados sobre maturidade, alocação e qualidade do trabalho
        * Forneça insights baseados em estatísticas e tendências
        * Use tom profissional para gestores
        * Sugira planos de ação e monitoramento
        * Cite dados específicos das análises fornecidas
        * Responda perguntas de follow-up mantendo o contexto da análise anterior
        * Quando houver gaps significativos entre métricas (ex: lead time vs cycle time), 
          explique possíveis causas e sugira investigações adicionais
        """
    }
    
    # Histórico de análise atual
    analise_atual = None
    entidade_atual = None
    
    print("Chat IA iniciado! Pergunte sobre tribos, squads ou peça insights.")
    print("Digite 'salvar' para exportar o chat para DOCX ou 'sair' para encerrar.")
    
    while True:
        query = input("Sua consulta: ").strip()
        
        if query.lower() == 'sair':
            print("\nEncerrando Agente Insights...")
            return 0
            
        if query.lower() == 'salvar':
            salvar_chat_docx(chat_log)
            continue
            
        try:
            # Se não houver análise atual ou a pergunta for sobre uma nova entidade
            if not analise_atual or "tribo" in query.lower() or "squad" in query.lower():
                # Identificar entidade na consulta
                entidade = identificar_entidade_consulta(query, analises)
                if not entidade:
                    print("Não foi possível identificar uma entidade específica na sua consulta.")
                    if analise_atual:
                        print("Continuando análise da entidade atual ou digite uma nova entidade.")
                    else:
                        tribos_disponiveis = list(estrutura['tribos'].keys()) if estrutura and 'tribos' in estrutura else []
                        squads_disponiveis = list(estrutura['squads'].keys()) if estrutura and 'squads' in estrutura else []
                        if tribos_disponiveis:
                            print(f"Tribos disponíveis: {', '.join(tribos_disponiveis)}")
                        if squads_disponiveis:
                            print(f"Squads disponíveis: {', '.join(squads_disponiveis)}")
                    continue
                    
                # Preparar dados para consulta
                dados_consulta = preparar_dados_consulta(entidade, entidade['nome'], analises)
                if not dados_consulta:
                    print("Não foi possível preparar os dados para análise.")
                    continue
                
                # Gerar nova análise
                analise_atual = gerar_analise_consultiva(dados_consulta.get('metricas', {}), 
                                                       dados_consulta.get('estrutura', {}),
                                                       {'tipo_entidade': entidade['tipo'],
                                                        'nome_entidade': entidade['nome'],
                                                        'query': query})
                entidade_atual = entidade
            
            # Gerar resposta contextualizada
            resposta = gerar_resposta_contextualizada(query, entidade_atual, dados_consulta, client)
            
            if resposta:
                print("\nResposta:")
                print(resposta)
                
                # Adiciona contexto para follow-up
                contexto_followup = {
                    "role": "system",
                    "content": f"""
                    Análise atual da {entidade_atual['tipo']} {entidade_atual['nome']}:
                    {formatar_analise_consultiva(analise_atual)}
                    
                    Use estas informações para responder perguntas de follow-up.
                    Se a pergunta for sobre gaps entre métricas, explique possíveis causas
                    e sugira investigações adicionais.
                    """
                }
                
                # Atualiza o contexto base com a análise atual
                contexto_base["content"] += f"\n\nContexto atual: {contexto_followup['content']}"
            
        except Exception as e:
            logging.error(f"Erro ao processar consulta: {str(e)}")
            print(f"\nErro ao processar consulta: {str(e)}")
            continue

def salvar_chat_docx(chat_log: List[tuple]):
    doc = Document()
    doc.add_heading('Chat de Insights - Agente Insights', 0)
    for autor, msg in chat_log:
        doc.add_paragraph(f"{autor}:", style='Heading 2')
        doc.add_paragraph(msg)
    caminho = f"output/relatorios/chat_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(caminho)
    print(f"Chat salvo em: {caminho}")

def analisar_alocacao(dados: pd.DataFrame, tribo: str = None, squad: str = None) -> Dict:
    """Analisa alocação de pessoas e papéis"""
    try:
        # Criar cópia dos dados
        df = dados.copy() if not dados.empty else pd.DataFrame()
        
        # Verificar colunas necessárias base (sem coluna de percentual)
        colunas_necessarias_base = ['endDate', 'role', 'squad', 'tribe', 'person']
        
        # Verificar variações da coluna de alocacao percentual
        coluna_percentual = None
        for possivel_coluna in ['percentageAllocation', 'percetageAllocation', 'percentage', 'alocacao_percentual']:
            if possivel_coluna in df.columns:
                coluna_percentual = possivel_coluna
                break
        
        # Log da coluna de percentual encontrada
        if coluna_percentual:
            logging.info(f"Coluna de percentual encontrada: {coluna_percentual}")
        else:
            logging.warning("Nenhuma coluna de percentual encontrada")
            
        # Verificar se temos as colunas base necessárias
        if not all(col in df.columns for col in colunas_necessarias_base):
            colunas_faltando = [col for col in colunas_necessarias_base if col not in df.columns]
            logging.warning(f"Colunas base ausentes para análise de alocação: {colunas_faltando}")
            return {
                'papeis': {},
                'alocacao_media': {},
                'pessoas_multi_squad': [],
                'composicao_squads': {},
                'media_pessoas_squad': 0,
                'pessoas_ativas': [],
                'squads_ativos': [],
                'tribos_ativas': [],
                'alocacoes_ativas': []
            }
        
        # Filtrar apenas alocações ativas (sem data de término ou data futura)
        df = df[df['endDate'].isna() | (pd.to_datetime(df['endDate'], errors='coerce') > pd.Timestamp.now())]
        logging.info(f"Alocações ativas encontradas: {len(df)}")
        if not df.empty:
            logging.info(f"Exemplo de tribos nas alocações: {df['tribe'].dropna().unique()[:5]}")
        
        # Filtrar por tribo ou squad se especificado
        if tribo:
            df = df[df['tribe'] == tribo]
        if squad:
            df = df[df['squad'] == squad]
            
        # Log dos registros após filtros
        logging.info(f"Registros após filtros de tribo/squad: {len(df)}")

        # Preparar dados para análise
        composicao_squads_agg = {}
        
        # Base para análise sem percentual
        composicao_squads_agg['role'] = lambda x: dict(Counter(x))
        
        # Adicionar percentual se disponível, com conversão segura para numérico
        if coluna_percentual:
            # Converter percentuais para valores numéricos com segurança
            try:
                # Se for string com %, remover e converter para float
                if df[coluna_percentual].dtype == 'object':
                    logging.info(f"Processando coluna {coluna_percentual}")
                    # Mostrar algumas amostras para debug
                    sample_values = df[coluna_percentual].dropna().head(5).tolist()
                    logging.info(f"Amostras de valores antes da conversão: {sample_values}")
                    
                    # Limpar e converter valores
                    df[coluna_percentual] = df[coluna_percentual].astype(str).str.replace('%', '').str.replace(',', '.')
                    df[coluna_percentual] = pd.to_numeric(df[coluna_percentual], errors='coerce')
                    
                    # Mostrar valores após limpeza para debug
                    clean_values = df[coluna_percentual].dropna().head(3).tolist()
                    logging.info(f"Valores após limpeza: {clean_values}")
                    
                    # Verificar valores inválidos (que viraram NaN)
                    invalid_count = df[coluna_percentual].isna().sum()
                    logging.info(f"Dados inválidos em {coluna_percentual}: {invalid_count} registros")
                
                # Calcular estatísticas básicas
                if df[coluna_percentual].notna().any():
                    min_val = df[coluna_percentual].min()
                    max_val = df[coluna_percentual].max()
                    mean_val = df[coluna_percentual].mean()
                    logging.info(f"Valores de percentual: min={min_val}, max={max_val}, mean={mean_val}")
                    
                    # Converter para decimal (0-1) se os valores parecem ser percentuais (>1)
                    if max_val > 1:
                        logging.info(f"Convertendo percentuais para decimal (0-1)")
                        df[coluna_percentual] = df[coluna_percentual] / 100
                        logging.info(f"Valores após normalização: min={df[coluna_percentual].min()}, max={df[coluna_percentual].max()}, mean={df[coluna_percentual].mean()}")
                
                # Usar mean para agregação apenas se os valores são numéricos
                composicao_squads_agg[coluna_percentual] = 'mean'
                
            except Exception as e:
                logging.warning(f"Erro ao processar coluna percentual {coluna_percentual}: {str(e)}")
                # Remover a coluna problemática
                if coluna_percentual in df.columns:
                    df = df.drop(coluna_percentual, axis=1)
                coluna_percentual = None
                
        # Calcular métricas de alocação
        analise = {
            'papeis': df.groupby('role').size().to_dict() if len(df) > 0 else {},
            'pessoas_multi_squad': df[df.groupby('person')['squad'].transform('size') > 1]['person'].unique().tolist() if len(df) > 0 else [],
            'media_pessoas_squad': float(df.groupby('squad').size().mean()) if len(df) > 0 and 'squad' in df.columns else 0,
            'pessoas_ativas': df['person'].dropna().unique().tolist() if 'person' in df.columns else [],
            'squads_ativos': df['squad'].dropna().unique().tolist() if 'squad' in df.columns else [],
            'tribos_ativas': df['tribe'].dropna().unique().tolist() if 'tribe' in df.columns else [],
            # Nova lista detalhada de alocações ativas
            'alocacoes_ativas': [
                {'pessoa': row['person'], 'squad': row['squad'], 'tribo': row['tribe']}
                for _, row in df.iterrows() if pd.notna(row['person']) and pd.notna(row['squad']) and pd.notna(row['tribe'])
            ]
        }
            
        # Adicionar métricas de alocação percentual se disponível
        if coluna_percentual and coluna_percentual in df.columns:
            analise['alocacao_media'] = df.groupby('squad')[coluna_percentual].mean().to_dict() if len(df) > 0 else {}
            
            # Usar agregação segura para composição de squads
            try:
                analise['composicao_squads'] = df.groupby('squad').agg(composicao_squads_agg).to_dict() if len(df) > 0 else {}
            except Exception as e:
                logging.warning(f"Erro na agregação de composição de squads: {str(e)}")
                # Fallback para composição sem percentual
                analise['composicao_squads'] = df.groupby('squad').agg({'role': lambda x: dict(Counter(x))}).to_dict() if len(df) > 0 else {}
        else:
            analise['alocacao_media'] = {}
            analise['composicao_squads'] = df.groupby('squad').agg({'role': lambda x: dict(Counter(x))}).to_dict() if len(df) > 0 else {}
        
        return analise
        
    except Exception as e:
        logging.error(f"Erro ao analisar alocação: {str(e)}")
        traceback.print_exc()
        return {
            'papeis': {},
            'alocacao_media': {},
            'pessoas_multi_squad': [],
            'composicao_squads': {},
            'media_pessoas_squad': 0,
            'pessoas_ativas': [],
            'squads_ativos': [],
            'tribos_ativas': [],
            'alocacoes_ativas': []
        }

def normalizar_texto(texto):
    """Normaliza um texto removendo acentos, convertendo para minúsculas e removendo caracteres especiais"""
    if not texto:
        return ""
    texto = str(texto).lower().strip()
    # Remove acentos
    nfkd = unicodedata.normalize('NFKD', texto)
    texto_sem_acentos = u"".join([c for c in nfkd if not unicodedata.combining(c)])
    # Remove caracteres especiais
    texto_limpo = re.sub(r'[^a-zA-Z0-9\s]', '', texto_sem_acentos)
    return texto_limpo.strip()

def analisar_cfd(df, colunas):
    """Analisa o Cumulative Flow Diagram do dataframe"""
    try:
        if 'date' not in colunas or 'status' not in colunas:
            return None
            
        # Converte coluna de data para datetime se necessário
        if not pd.api.types.is_datetime64_any_dtype(df[colunas['date']]):
            df[colunas['date']] = pd.to_datetime(df[colunas['date']], errors='coerce')
            
        # Remove linhas com datas inválidas
        df = df.dropna(subset=[colunas['date']])
        
        # Ordena por data
        df = df.sort_values(by=colunas['date'])
        
        # Calcula métricas
        total_items = len(df)
        dias_total = (df[colunas['date']].max() - df[colunas['date']].min()).days
        
        if dias_total == 0:
            return {
                'throughput': total_items,
                'throughput_diario': total_items,
                'avg_lead_time': 0,
                'wip': total_items
            }
        
        # Calcula throughput
        throughput = total_items
        throughput_diario = throughput / max(dias_total, 1)
        
        # Calcula WIP (Work in Progress)
        status_em_andamento = ['in progress', 'em andamento', 'doing']
        wip = len(df[df[colunas['status']].str.lower().isin(status_em_andamento)])
        
        # Calcula lead time médio
        status_concluido = ['done', 'concluido', 'completed', 'finalizado']
        df_concluidos = df[df[colunas['status']].str.lower().isin(status_concluido)]
        
        if len(df_concluidos) > 0:
            lead_times = []
            for _, item in df_concluidos.iterrows():
                created_date = item[colunas['date']]
                done_date = item[colunas['date']]  # Assumindo que a data é a de conclusão
                lead_time = (done_date - created_date).days
                lead_times.append(max(0, lead_time))
            avg_lead_time = sum(lead_times) / len(lead_times) if lead_times else 0
        else:
            avg_lead_time = 0
            
        return {
            'throughput': throughput,
            'throughput_diario': throughput_diario,
            'avg_lead_time': avg_lead_time,
            'wip': wip
        }
        
    except Exception as e:
        logging.error(f"Erro ao analisar CFD: {str(e)}")
        return None

def mapear_estrutura_org(analises):
    """
    Mapeia a estrutura organizacional a partir das análises, considerando apenas alocações ativas (endDate em branco ou futura) e tribos presentes no arquivo de maturidade.
    """
    estrutura = {
        'tribos': {},
        'squads': {},
        'pessoas': set(),
        'total_squads': 0,
        'total_pessoas': 0,
        'papeis_total': {}
    }
    squads_unicos = set()
    pessoas_unicas = set()
    papeis = defaultdict(int)
    tribos_maturidade = set()
    # Buscar tribos apenas da análise de maturidade
    for analise in analises:
        if isinstance(analise, dict) and analise.get('fonte') == 'maturidade' and 'tribos_ativas' in analise:
            tribos_maturidade.update(analise['tribos_ativas'])
    logging.info(f"Tribos de maturidade: {list(tribos_maturidade)}")
    # Inicializa estrutura das tribos válidas
    estrutura['tribos'] = {tribo: {'squads': set(), 'pessoas': set()} for tribo in tribos_maturidade}
    # Preencher squads e pessoas de cada tribo válida a partir das alocações ativas
    total_matches = 0
    for analise in analises:
        if not isinstance(analise, dict):
            continue
        alocacoes = analise.get('alocacoes_ativas', [])
        for aloc in alocacoes:
            tribo = aloc['tribo']
            squad = aloc['squad']
            pessoa = aloc['pessoa']
            if tribo in estrutura['tribos']:
                estrutura['tribos'][tribo]['squads'].add(squad)
                estrutura['tribos'][tribo]['pessoas'].add(pessoa)
                squads_unicos.add(squad)
                pessoas_unicas.add(pessoa)
                total_matches += 1
    logging.info(f"Total de matches tribo alocação x maturidade: {total_matches}")
    logging.info(f"Exemplo de tribos válidas na estrutura: {list(estrutura['tribos'].keys())[:5]}")
    estrutura['squads'] = {squad: {} for squad in squads_unicos}
    estrutura['pessoas'] = list(pessoas_unicas)
    estrutura['total_squads'] = len(squads_unicos)
    estrutura['total_pessoas'] = len(pessoas_unicas)
    # Converte sets para listas
    for tribo in estrutura['tribos'].values():
        tribo['squads'] = list(tribo['squads'])
        tribo['pessoas'] = list(tribo['pessoas'])
    # Processa papéis totais se disponível
    for analise in analises:
        if 'papeis' in analise:
            for papel, qtd in analise['papeis'].items():
                papeis[papel] += qtd
    estrutura['papeis_total'] = dict(papeis)
    return estrutura

def identificar_entidade_consulta(query: str, analises: List[Dict]) -> Optional[Dict]:
    """
    Identifica a entidade (tribo ou squad) mencionada na consulta do usuário.
    Retorna um dicionário com tipo ('tribo' ou 'squad') e nome, ou None se não encontrar.
    Se não encontrar correspondência clara, retorna as opções mais próximas para o usuário escolher.
    """
    def normalizar(texto):
        return unidecode(str(texto)).lower().replace(' ', '').replace('-', '').replace('&', '').replace('/', '').replace('_', '')

    query_norm = normalizar(query)
    nomes_tribos = [a['nome'] for a in analises if a.get('tipo') == 'tribo']
    nomes_squads = [a['nome'] for a in analises if a.get('tipo') == 'squad']
    nomes_tribos_norm = [normalizar(n) for n in nomes_tribos]
    nomes_squads_norm = [normalizar(n) for n in nomes_squads]

    # 1. Correspondência exata
    if query_norm in nomes_tribos_norm:
        idx = nomes_tribos_norm.index(query_norm)
        return {'tipo': 'tribo', 'nome': nomes_tribos[idx]}
    if query_norm in nomes_squads_norm:
        idx = nomes_squads_norm.index(query_norm)
        return {'tipo': 'squad', 'nome': nomes_squads[idx]}

    # 2. Correspondência parcial
    for i, n in enumerate(nomes_tribos_norm):
        if query_norm in n or n in query_norm:
            return {'tipo': 'tribo', 'nome': nomes_tribos[i]}
    for i, n in enumerate(nomes_squads_norm):
        if query_norm in n or n in query_norm:
            return {'tipo': 'squad', 'nome': nomes_squads[i]}

    # 3. Fuzzy matching (difflib)
    tribo_match = difflib.get_close_matches(query_norm, nomes_tribos_norm, n=3, cutoff=0.6)
    squad_match = difflib.get_close_matches(query_norm, nomes_squads_norm, n=3, cutoff=0.6)
    sugestoes = []
    if tribo_match:
        sugestoes.extend([nomes_tribos[nomes_tribos_norm.index(m)] for m in tribo_match])
    if squad_match:
        sugestoes.extend([nomes_squads[nomes_squads_norm.index(m)] for m in squad_match])
    if sugestoes:
        print("Não foi possível identificar exatamente a entidade. Você quis dizer uma destas opções?")
        for i, s in enumerate(sugestoes, 1):
            print(f"{i}. {s}")
        print("Digite o número da opção desejada ou tente novamente.")
        escolha = input("Opção: ").strip()
        if escolha.isdigit():
            idx = int(escolha) - 1
            if 0 <= idx < len(sugestoes):
                nome_escolhido = sugestoes[idx]
                if nome_escolhido in nomes_tribos:
                    return {'tipo': 'tribo', 'nome': nome_escolhido}
                if nome_escolhido in nomes_squads:
                    return {'tipo': 'squad', 'nome': nome_escolhido}
        return None
    # 4. Não encontrou nada
    return None

def preparar_dados_consulta(entidade: Dict, nome: str, analises: List[Dict]) -> Optional[Dict]:
    """Prepara os dados relevantes para a consulta"""
    try:
        # Encontra a análise específica para a entidade
        analise_entidade = None
        for analise in analises:
            if analise.get('tipo') == entidade['tipo'] and analise.get('nome') == nome:
                analise_entidade = analise
                break
        
        if not analise_entidade:
            return None
            
        # Extrai os insights da análise
        insights = analise_entidade.get('insights', {})
        
        # Prepara os dados no formato esperado
        dados = {
            'metricas': {
                'lead_time': {
                    'medio': insights.get('lead_time_medio', 0),
                    'mediana': insights.get('lead_time_mediana', 0),
                    'p75': insights.get('lead_time_p75', 0),
                    'p95': insights.get('lead_time_p95', 0)
                },
                'cycle_time': {
                    'medio': insights.get('cycle_time_medio', 0),
                    'mediana': insights.get('cycle_time_mediana', 0),
                    'p75': insights.get('cycle_time_p75', 0),
                    'p95': insights.get('cycle_time_p95', 0)
                },
                'throughput': insights.get('throughput', 0),
                'story_points': insights.get('story_points_medio', 0)
            },
            'analises_complementares': {
                'total_pessoas': insights.get('total_pessoas', 0),
                'total_squads': insights.get('total_squads', 0) if entidade['tipo'] == 'tribo' else insights.get('total_tribos', 0)
            },
            'metricas_ageis': {
                'lead_time': {
                    'avg': insights.get('lead_time_medio', 0),
                    'min': insights.get('lead_time_mediana', 0),
                    'max': insights.get('lead_time_p95', 0),
                    'tendencia': 'estável' if insights.get('lead_time_medio', 0) < 30 else 'crescente'
                },
                'cycle_time': {
                    'avg': insights.get('cycle_time_medio', 0),
                    'min': insights.get('cycle_time_mediana', 0),
                    'max': insights.get('cycle_time_p95', 0)
                },
                'throughput': {
                    'atual': insights.get('throughput', 0),
                    'tendencia': 'estável' if insights.get('throughput', 0) > 100 else 'crescente'
                }
            },
            'composicao_time': {
                'total_pessoas': insights.get('total_pessoas', 0),
                'total_squads': insights.get('total_squads', 0) if entidade['tipo'] == 'tribo' else insights.get('total_tribos', 0)
            },
            'contexto': {
                'tipo': entidade['tipo'],
                'nome': nome,
                'maturidade': insights.get('maturidade', 0) if 'maturidade' in insights else None
            },
            'analises': [{
                'tipo': 'analise_consultiva',
                'dados': {
                    'diagnostico': {
                        'lead_time': insights.get('lead_time_medio', 0),
                        'throughput': insights.get('throughput', 0),
                        'story_points': insights.get('story_points_medio', 0)
                    },
                    'recomendacoes': [],
                    'insights': []
                }
            }]
        }
        
        return dados
    except Exception as e:
        logging.error(f"Erro ao preparar dados: {str(e)}")
        return None

def analisar_tribo(tribo, estrutura, analises):
    """Análise aprofundada de uma tribo específica"""
    dados = {
        'metricas_ageis': {},
        'qualidade_entrega': {},
        'eficiencia_fluxo': {},
        'capacidade_times': {},
        'tendencias': {},
        'insights_descritivos': []
    }
    for analise in analises:
        if not isinstance(analise, dict):
            continue
        # Análise de maturidade e eficiência
        if 'metricas_por_tribo' in analise and tribo in analise['metricas_por_tribo']:
            metricas = analise['metricas_por_tribo'][tribo]
            dados['metricas_ageis'] = extrair_metricas_ageis(metricas)
            dados['qualidade_entrega'] = analisar_qualidade_entrega(metricas)
            dados['insights_descritivos'] = gerar_insights_descritivos(dados['metricas_ageis'])
        # Análise de pessoas e capacidade
        if 'composicao_squads' in analise:
            dados['capacidade_times'] = analisar_capacidade_times(
                analise['composicao_squads'],
                estrutura['tribos'][tribo]['squads']
            )
    return dados

def analisar_squad(squad, estrutura, analises):
    """Análise aprofundada de um squad específico"""
    dados = {
        'metricas_ageis': {},
        'qualidade_entrega': {},
        'eficiencia_fluxo': {},
        'composicao_time': {},
        'tendencias': {}
    }
    
    for analise in analises:
        if not isinstance(analise, dict):
            continue
            
        # Métricas do squad
        if 'metricas_por_squad' in analise and squad in analise['metricas_por_squad']:
            metricas = analise['metricas_por_squad'][squad]
            dados['metricas_ageis'] = extrair_metricas_ageis(metricas)
            dados['qualidade_entrega'] = analisar_qualidade_entrega(metricas)
            
        # Análise de composição e capacidade
        if 'composicao_squads' in analise and squad in analise['composicao_squads']:
            dados['composicao_time'] = analise['composicao_squads'][squad]
            dados['eficiencia_fluxo'] = analisar_eficiencia_fluxo(
                metricas,
                analise['composicao_squads'][squad]
            )
    
    return dados

def extrair_metricas_ageis(metricas):
    """Extrai e analisa métricas ágeis"""
    return {
        'lead_time': calcular_metricas_lead_time(metricas),
        'cycle_time': calcular_metricas_cycle_time(metricas),
        'throughput': calcular_metricas_throughput(metricas),
        'wip': calcular_metricas_wip(metricas)
    }

def calcular_metricas_lead_time(metricas):
    """Calcula métricas de lead time com análise de tendências"""
    lead_times = []
    if 'items' in metricas:
        for item in metricas['items']:
            if 'created_date' in item and 'done_date' in item:
                lead_time = (item['done_date'] - item['created_date']).days
                lead_times.append(lead_time)
    
    if not lead_times:
        return {'avg': 0, 'min': 0, 'max': 0, 'tendencia': 'estável'}
        
    return {
        'avg': np.mean(lead_times),
        'min': np.min(lead_times),
        'max': np.max(lead_times),
        'tendencia': analisar_tendencia(lead_times)
    }

def calcular_metricas_cycle_time(metricas):
    """Calcula métricas de cycle time com análise de gargalos"""
    cycle_times = []
    tempos_por_status = defaultdict(list)
    
    if 'items' in metricas:
        for item in metricas['items']:
            if 'status_changes' in item:
                cycle_time = 0
                for status, tempo in item['status_changes'].items():
                    cycle_time += tempo
                    tempos_por_status[status].append(tempo)
                cycle_times.append(cycle_time)
    
    gargalos = identificar_gargalos(tempos_por_status)
    
    return {
        'avg': np.mean(cycle_times) if cycle_times else 0,
        'gargalos': gargalos,
        'distribuicao': calcular_distribuicao_tempos(cycle_times)
    }

def calcular_metricas_throughput(metricas):
    """Calcula throughput com previsões"""
    entregas_por_periodo = []
    if 'items' in metricas:
        # Agrupa entregas por período
        periodos = agrupar_entregas_por_periodo(metricas['items'])
        entregas_por_periodo = list(periodos.values())
    
    return {
        'atual': np.mean(entregas_por_periodo[-4:]) if len(entregas_por_periodo) >= 4 else 0,
        'tendencia': analisar_tendencia(entregas_por_periodo),
        'previsao': prever_entregas(entregas_por_periodo) if entregas_por_periodo else 0,
        'estabilidade': calcular_estabilidade(entregas_por_periodo)
    }

def agrupar_entregas_por_periodo(items):
    """Agrupa entregas por período (mês)"""
    periodos = defaultdict(int)
    for item in items:
        if 'data_conclusao' in item:
            data = datetime.strptime(item['data_conclusao'], '%Y-%m-%d')
            periodo = f"{data.year}-{data.month:02d}"
            periodos[periodo] += 1
    return dict(sorted(periodos.items()))

def analisar_tendencia(dados):
    """Analisa tendência dos dados usando regressão linear"""
    if not dados or len(dados) < 2:
        return 'estável'
    
    x = np.arange(len(dados))
    y = np.array(dados)
    slope, _ = np.polyfit(x, y, 1)
    
    if abs(slope) < 0.1:
        return 'estável'
    return 'crescente' if slope > 0 else 'decrescente'

def prever_entregas(dados, periodos_futuros=3):
    """Prevê entregas para os próximos períodos usando média móvel"""
    if not dados or len(dados) < 4:
        return 0
    
    media_movel = np.mean(dados[-4:])
    return round(media_movel)

def calcular_estabilidade(dados):
    """Calcula estabilidade do fluxo usando coeficiente de variação"""
    if not dados or len(dados) < 2:
        return 'não há dados suficientes'
    
    cv = np.std(dados) / np.mean(dados) if np.mean(dados) > 0 else float('inf')
    
    if cv < 0.2:
        return 'muito estável'
    elif cv < 0.4:
        return 'estável'
    elif cv < 0.6:
        return 'moderadamente instável'
    else:
        return 'instável'

def identificar_gargalos(tempos_por_status):
    """Identifica gargalos no processo baseado nos tempos por status"""
    gargalos = []
    if not tempos_por_status:
        return gargalos
        
    tempo_total_medio = sum(np.mean(tempos) for tempos in tempos_por_status.values())
    
    for status, tempos in tempos_por_status.items():
        tempo_medio = np.mean(tempos)
        if tempo_medio > tempo_total_medio * 0.3:  # Status que consomem mais de 30% do tempo total
            gargalos.append({
                'status': status,
                'tempo_medio': tempo_medio,
                'percentual': (tempo_medio / tempo_total_medio) * 100
            })
    
    return sorted(gargalos, key=lambda x: x['tempo_medio'], reverse=True)

def calcular_distribuicao_tempos(tempos):
    """Calcula distribuição dos tempos para análise estatística"""
    if not tempos:
        return {
            'min': 0,
            'max': 0,
            'p25': 0,
            'p50': 0,
            'p75': 0
        }
        
    return {
        'min': np.min(tempos),
        'max': np.max(tempos),
        'p25': np.percentile(tempos, 25),
        'p50': np.percentile(tempos, 50),
        'p75': np.percentile(tempos, 75)
    }

def calcular_metricas_wip(metricas):
    """Calcula métricas de Work in Progress (WIP)"""
    wip_atual = 0
    wip_por_status = defaultdict(int)
    wip_historico = []
    
    if 'items' in metricas:
        # Calcula WIP atual
        for item in metricas['items']:
            if item.get('status') != 'Concluído' and item.get('status') != 'Cancelado':
                wip_atual += 1
                wip_por_status[item.get('status', 'Desconhecido')] += 1
        
        # Calcula histórico de WIP (se disponível)
        if 'wip_historico' in metricas:
            wip_historico = metricas['wip_historico']
    
    return {
        'atual': wip_atual,
        'por_status': dict(wip_por_status),
        'tendencia': analisar_tendencia(wip_historico) if wip_historico else 'sem dados históricos',
        'limite_recomendado': calcular_limite_wip(wip_historico) if wip_historico else None
    }

def calcular_limite_wip(historico):
    """Calcula limite recomendado de WIP baseado no histórico"""
    if not historico or len(historico) < 5:
        return None
        
    # Usa percentil 85 como limite superior recomendado
    return np.percentile(historico, 85)

def analisar_capacidade_times(metricas, estrutura):
    """Analisa capacidade dos times baseado em métricas históricas"""
    capacidade = {
        'media_entregas_mes': 0,
        'capacidade_atual': 'não há dados suficientes',
        'sugestoes': []
    }
    
    if 'items' in metricas and len(metricas['items']) > 0:
        # Calcula média de entregas por mês
        entregas_por_periodo = []
        periodos = agrupar_entregas_por_periodo(metricas['items'])
        entregas_por_periodo = list(periodos.values())
        
        if entregas_por_periodo:
            media_entregas = np.mean(entregas_por_periodo[-3:]) if len(entregas_por_periodo) >= 3 else np.mean(entregas_por_periodo)
            capacidade['media_entregas_mes'] = round(media_entregas, 1)
            
            # Avalia capacidade
            if media_entregas < 2:
                capacidade['capacidade_atual'] = 'abaixo do esperado'
                capacidade['sugestoes'].append('Considerar redistribuição de recursos ou identificar impedimentos')
            elif media_entregas < 4:
                capacidade['capacidade_atual'] = 'adequada'
            else:
                capacidade['capacidade_atual'] = 'acima da média'
                
    return capacidade

def analisar_qualidade_entrega(metricas):
    """Analisa qualidade das entregas baseado em retrabalho e bugs"""
    qualidade = {
        'taxa_retrabalho': 0,
        'taxa_bugs': 0,
        'nivel_qualidade': 'não há dados suficientes',
        'sugestoes': []
    }
    
    if 'items' in metricas and len(metricas['items']) > 0:
        total_items = len(metricas['items'])
        retrabalho = sum(1 for item in metricas['items'] if item.get('retrabalho', False))
        bugs = sum(1 for item in metricas['items'] if item.get('tipo') == 'bug')
        
        qualidade['taxa_retrabalho'] = round((retrabalho / total_items) * 100, 1)
        qualidade['taxa_bugs'] = round((bugs / total_items) * 100, 1)
        
        # Avalia qualidade
        if qualidade['taxa_retrabalho'] > 20 or qualidade['taxa_bugs'] > 15:
            qualidade['nivel_qualidade'] = 'precisa melhorar'
            qualidade['sugestoes'].append('Aumentar cobertura de testes')
            qualidade['sugestoes'].append('Revisar processo de code review')
        elif qualidade['taxa_retrabalho'] > 10 or qualidade['taxa_bugs'] > 7:
            qualidade['nivel_qualidade'] = 'moderada'
            qualidade['sugestoes'].append('Considerar melhorias no processo de QA')
        else:
            qualidade['nivel_qualidade'] = 'boa'
            
    return qualidade

def analisar_eficiencia_fluxo(metricas):
    """Analisa eficiência do fluxo de trabalho"""
    eficiencia = {
        'tempo_medio_conclusao': 0,
        'taxa_bloqueio': 0,
        'nivel_eficiencia': 'não há dados suficientes',
        'sugestoes': []
    }
    
    if 'items' in metricas and len(metricas['items']) > 0:
        tempos_conclusao = []
        total_bloqueios = 0
        
        for item in metricas['items']:
            if 'tempo_total' in item:
                tempos_conclusao.append(item['tempo_total'])
            if item.get('teve_bloqueio', False):
                total_bloqueios += 1
                
        if tempos_conclusao:
            eficiencia['tempo_medio_conclusao'] = round(np.mean(tempos_conclusao), 1)
            eficiencia['taxa_bloqueio'] = round((total_bloqueios / len(metricas['items'])) * 100, 1)
            
            # Avalia eficiência
            if eficiencia['taxa_bloqueio'] > 30:
                eficiencia['nivel_eficiencia'] = 'baixa'
                eficiencia['sugestoes'].append('Investigar causas frequentes de bloqueio')
            elif eficiencia['taxa_bloqueio'] > 15:
                eficiencia['nivel_eficiencia'] = 'moderada'
                eficiencia['sugestoes'].append('Monitorar impedimentos recorrentes')
            else:
                eficiencia['nivel_eficiencia'] = 'boa'
                
    return eficiencia

def preparar_analise_geral(estrutura, analises):
    """Prepara análise geral com insights principais"""
    insight_geral = {
        'resumo': 'Análise geral do desempenho',
        'pontos_positivos': [],
        'pontos_atencao': [],
        'recomendacoes': []
    }
    
    # Analisa throughput geral
    throughputs = []
    for analise in analises:
        if isinstance(analise, dict) and 'metricas' in analise:
            throughput = analise['metricas'].get('throughput', {}).get('atual', 0)
            if throughput > 0:
                throughputs.append(throughput)
                
    if throughputs:
        media_throughput = np.mean(throughputs)
        if media_throughput > 5:
            insight_geral['pontos_positivos'].append('Boa taxa de entrega geral')
        elif media_throughput < 2:
            insight_geral['pontos_atencao'].append('Taxa de entrega abaixo do esperado')
            insight_geral['recomendacoes'].append('Avaliar capacidade e impedimentos das equipes')
    
    # Analisa lead times
    lead_times = []
    for analise in analises:
        if isinstance(analise, dict) and 'metricas' in analise:
            lt = analise['metricas'].get('lead_time', {}).get('avg', 0)
            if lt > 0:
                lead_times.append(lt)
    
    if lead_times:
        media_lead_time = np.mean(lead_times)
        if media_lead_time > 20:
            insight_geral['pontos_atencao'].append('Lead time elevado')
            insight_geral['recomendacoes'].append('Identificar e eliminar gargalos no processo')
        elif media_lead_time < 10:
            insight_geral['pontos_positivos'].append('Lead time dentro do esperado')
    
    # Adiciona recomendações gerais se necessário
    if not insight_geral['pontos_positivos'] and not insight_geral['pontos_atencao']:
        insight_geral['recomendacoes'].append('Coletar mais dados para análise aprofundada')
    
    return insight_geral

def analisar_metricas_complementares(entidade, nome, analises):
    """Analisa métricas complementares para uma entidade específica"""
    metricas_comp = {}
    
    if entidade in analises and nome in analises[entidade]:
        dados = analises[entidade][nome]
        if 'metricas' in dados:
            metricas = dados['metricas']
            estrutura = dados.get('estrutura', {})
            
            # Análise de capacidade
            metricas_comp['capacidade_times'] = analisar_capacidade_times(metricas, estrutura)
            
            # Análise de qualidade
            metricas_comp['qualidade_entrega'] = analisar_qualidade_entrega(metricas)
            
            # Análise de eficiência
            metricas_comp['eficiencia_fluxo'] = analisar_eficiencia_fluxo(metricas)
    
    return metricas_comp

def gerar_resposta_contextualizada(query: str, entidade: Dict, dados: Dict, client: Any) -> str:
    """
    Gera uma resposta contextualizada baseada na consulta e dados disponíveis.
    Mantém o contexto da análise anterior para perguntas de follow-up.
    """
    try:
        # Verifica se há dados suficientes
        if not dados or not dados.get('metricas'):
            return "Não há dados suficientes para gerar uma análise."
            
        # Extrai métricas principais
        metricas = dados['metricas']
        estrutura = dados.get('estrutura', {})
        
        # Prepara o prompt para análise de gaps
        prompt_gaps = """
        Analise os seguintes gaps e forneça insights:
        1. Lead Time vs Cycle Time:
           - Lead Time médio: {lead_time_medio} dias
           - Cycle Time médio: {cycle_time_medio} dias
           - Gap: {gap_lt_ct} dias
        
        2. Possíveis causas do gap:
           - Tempo de espera antes do desenvolvimento
           - Tempo de revisão e aprovação
           - Tempo de deploy e homologação
           - Bloqueios e dependências
        
        3. Recomendações para investigação:
           - Análise de fluxo de trabalho
           - Mapeamento de gargalos
           - Otimização de processos
        """
        
        # Prepara o prompt para análise de tendências
        prompt_tendencias = """
        Analise as seguintes tendências:
        1. Throughput:
           - Total: {throughput_total}
           - Média mensal: {throughput_medio}
        
        2. Complexidade:
           - Story points médios: {story_points_medio}
           - Distribuição de complexidade
        
        3. Eficiência:
           - Lead Time vs Cycle Time
           - Taxa de entrega
           - Qualidade das entregas
        """
        
        # Prepara o prompt para análise de estrutura
        prompt_estrutura = """
        Analise a estrutura organizacional:
        1. Composição:
           - Total de pessoas: {total_pessoas}
           - Número de squads: {num_squads}
           - Média de pessoas por squad: {media_pessoas_squad}
        
        2. Distribuição:
           - Papéis e responsabilidades
           - Capacidade de entrega
           - Balanceamento de carga
        """
        
        # Gera a análise consultiva
        analise = gerar_analise_consultiva(metricas, estrutura, {
            'tipo_entidade': entidade['tipo'],
            'nome_entidade': entidade['nome'],
            'query': query
        })
        
        # Formata a análise
        resposta = formatar_analise_consultiva(analise)
        
        # Adiciona análise de gaps se relevante
        if 'lead_time' in metricas and 'cycle_time' in metricas:
            gap_lt_ct = metricas['lead_time']['media'] - metricas['cycle_time']['media']
            if gap_lt_ct > 5:  # Se o gap for significativo
                resposta += "\n\nAnálise de Gap (Lead Time vs Cycle Time):"
                resposta += f"\n- Gap identificado: {gap_lt_ct:.1f} dias"
                resposta += "\n- Possíveis causas:"
                resposta += "\n  * Tempo de espera antes do desenvolvimento"
                resposta += "\n  * Tempo de revisão e aprovação"
                resposta += "\n  * Tempo de deploy e homologação"
                resposta += "\n  * Bloqueios e dependências"
                resposta += "\n- Recomendações para investigação:"
                resposta += "\n  * Realizar análise detalhada do fluxo de trabalho"
                resposta += "\n  * Mapear gargalos no processo"
                resposta += "\n  * Identificar pontos de otimização"
        
        return resposta
        
    except Exception as e:
        logging.error(f"Erro ao gerar resposta contextualizada: {str(e)}")
        return f"Erro ao gerar análise: {str(e)}"

def get_device():
    """Detecta e retorna o dispositivo mais apropriado para processamento"""
    if torch.cuda.is_available():
        logging.info("GPU disponível: usando CUDA para aceleração")
        return "cuda"
    else:
        logging.info("GPU não disponível: usando CPU")
        return "cpu"

def to_device(data, device):
    """Converte dados para o dispositivo apropriado"""
    if device == "cuda":
        if isinstance(data, pd.DataFrame):
            # Converte DataFrame para GPU usando CuPy
            return cp.array(data.values)
        elif isinstance(data, np.ndarray):
            return cp.array(data)
        elif isinstance(data, torch.Tensor):
            return data.cuda()
    return data

def from_device(data, device):
    """Converte dados de volta do dispositivo para CPU"""
    if device == "cuda":
        if isinstance(data, cp.ndarray):
            return cp.asnumpy(data)
        elif isinstance(data, torch.Tensor):
            return data.cpu().numpy()
    return data

def processar_dados_paralelo(dados: Dict[str, pd.DataFrame], device: str) -> Dict[str, Any]:
    """Processa dados em paralelo usando GPU ou CPU"""
    with ThreadPoolExecutor() as executor:
        futures = []
        
        # Processa cada DataFrame em paralelo
        for nome, df in dados.items():
            future = executor.submit(processar_dataframe, df, device)
            futures.append((nome, future))
        
        # Coleta resultados
        resultados = {}
        for nome, future in futures:
            resultados[nome] = future.result()
            
    return resultados

def processar_dataframe(df: pd.DataFrame, device: str) -> pd.DataFrame:
    """Processa um DataFrame usando GPU ou CPU"""
    # Converte para o dispositivo apropriado
    dados = to_device(df, device)
    
    if device == "cuda":
        # Operações otimizadas para GPU
        # Exemplo: cálculos estatísticos
        if isinstance(dados, cp.ndarray):
            media = cp.mean(dados, axis=0)
            desvio = cp.std(dados, axis=0)
            # Converte de volta para CPU
            return pd.DataFrame({
                'media': from_device(media, device),
                'desvio': from_device(desvio, device)
            })
    else:
        # Operações CPU padrão
        return df

def executar_pipeline():
    """Executa o pipeline completo de análise com suporte a GPU"""
    try:
        # Detecta dispositivo
        device = get_device()
        logging.info(f"Usando dispositivo: {device}")
        
        # Carregar dados
        dados = carregar_dados()
        if not dados:
            logging.error("Falha ao carregar dados")
            return None
            
        # Processa dados em paralelo
        dados_processados = processar_dados_paralelo(dados, device)
        
        # Padronizar IDs
        dados = padronizar_ids(dados)
        if not dados:
            logging.error("Falha ao padronizar IDs")
            return None
            
        # Cruzar dados robusto
        if 'maturidade' in dados and 'alocacao' in dados and 'executivo' in dados:
            df_cruzado = cruzar_dados_completo(dados['maturidade'], dados['alocacao'], dados['executivo'])
            print(f"Cruzamento completo: {len(df_cruzado)} linhas")
            print(df_cruzado[['Tribo', 'tribe', 'squad', 'Ano', 'Quarter', 'Chave_DataTribo', 'Chave_DataSquad']].head())
            
            # Gerar estrutura e insights
            estrutura, insights_tribos, insights_squads = gerar_estrutura_e_insights(df_cruzado)
            print(f"Tribos reconhecidas: {list(estrutura['tribos'].keys())}")
            print(f"Squads reconhecidos: {list(estrutura['squads'].keys())[:10]}")
            
            # Exemplos de insights
            print(f"Exemplo de insights para tribo:")
            for tribo, insight in list(insights_tribos.items())[:3]:
                print(f"- {tribo}: {insight}")
            print(f"Exemplo de insights para squad:")
            for squad, insight in list(insights_squads.items())[:3]:
                print(f"- {squad}: {insight}")
        else:
            logging.error("Dados insuficientes para cruzamento completo")
            return None
            
        # Gera análises
        analises = []
        for tribo, insight in insights_tribos.items():
            analises.append({
                "tipo": "tribo",
                "nome": tribo,
                "insights": insight,
                "descricao": f"Análise da tribo {tribo} com {insight.get('total_pessoas', 0)} pessoas e {insight.get('total_squads', 0)} squads. Lead time médio: {insight.get('lead_time_medio', 0):.1f} dias, throughput: {insight.get('throughput', 0)} entregas."
            })
            
        for squad, insight in insights_squads.items():
            analises.append({
                "tipo": "squad",
                "nome": squad,
                "insights": insight,
                "descricao": f"Análise do squad {squad} com {insight.get('total_pessoas', 0)} pessoas. Lead time médio: {insight.get('lead_time_medio', 0):.1f} dias, throughput: {insight.get('throughput', 0)} entregas."
            })
            
        return analises
        
    except Exception as e:
        logging.error(f"Erro no pipeline de análise: {str(e)}")
        traceback.print_exc()
        return None

def gerar_insights_descritivos(metricas_tribo):
    """Gera insights descritivos baseados nas métricas reais da tribo."""
    insights = []
    # Lead time
    lt = metricas_tribo.get('lead_time', {})
    if lt and isinstance(lt, dict):
        insights.append(f"Lead time médio: {lt.get('avg', 0):.1f} dias (mín: {lt.get('min', 0)}, máx: {lt.get('max', 0)}). Tendência: {lt.get('tendencia', 'indefinida')}.")
        if lt.get('tendencia') == 'crescente':
            insights.append("Atenção: tendência de aumento no lead time, indicando possíveis gargalos recentes.")
        elif lt.get('tendencia') == 'decrescente':
            insights.append("Ótimo: tendência de redução no lead time, mostrando melhoria no fluxo.")
    # Throughput
    th = metricas_tribo.get('throughput', {})
    if th and isinstance(th, dict):
        insights.append(f"Throughput médio recente: {th.get('atual', 0):.1f} entregas/mês. Tendência: {th.get('tendencia', 'indefinida')}.")
        if th.get('tendencia') == 'decrescente':
            insights.append("Alerta: throughput em queda, pode indicar sobrecarga ou bloqueios.")
        elif th.get('tendencia') == 'crescente':
            insights.append("Ótimo: throughput em crescimento, indicando aumento de produtividade.")
    # Eficiência
    ef = metricas_tribo.get('eficiencia_fluxo', {})
    if ef and isinstance(ef, dict):
        insights.append(f"Eficiência do fluxo: {ef.get('nivel_eficiencia', 'indefinida')}. Tempo médio de conclusão: {ef.get('tempo_medio_conclusao', 0):.1f} dias.")
        if ef.get('nivel_eficiencia') == 'baixa':
            insights.append("Atenção: eficiência baixa, recomenda-se investigar causas de bloqueio.")
    # WIP
    wip = metricas_tribo.get('wip', {})
    if wip and isinstance(wip, dict):
        insights.append(f"WIP atual: {wip.get('atual', 0)} itens em andamento.")
    # Qualidade
    qual = metricas_tribo.get('qualidade_entrega', {})
    if qual and isinstance(qual, dict):
        insights.append(f"Taxa de retrabalho: {qual.get('taxa_retrabalho', 0):.1f}%. Taxa de bugs: {qual.get('taxa_bugs', 0):.1f}%. Nível: {qual.get('nivel_qualidade', 'indefinido')}.")
        if qual.get('nivel_qualidade') == 'precisa melhorar':
            insights.append("Alerta: qualidade abaixo do esperado, recomenda-se revisão do processo de QA.")
    # Pontos de atenção gerais
    if not insights:
        insights.append("Não há dados suficientes para gerar insights descritivos detalhados para esta tribo.")
    return insights

def gerar_analise_consultiva(metricas: Dict, estrutura: Dict, contexto: Dict) -> Dict:
    """
    Gera análise consultiva completa baseada em dados e métricas.
    Implementa análises descritivas, diagnósticas, preditivas e prescritivas.
    """
    analise = {
        'tipo': 'analise_consultiva',
        'timestamp': datetime.now().isoformat(),
        'metricas_principais': {},
        'diagnosticos': [],
        'previsoes': [],
        'recomendacoes': [],
        'riscos': [],
        'oportunidades': [],
        'planos_acao': []
    }
    
    # Análise Descritiva
    analise['metricas_principais'] = {
        'lead_time': {
            'medio': metricas.get('lead_time_medio', 0),
            'mediana': metricas.get('lead_time_mediana', 0),
            'p75': metricas.get('lead_time_p75', 0),
            'p95': metricas.get('lead_time_p95', 0)
        },
        'cycle_time': {
            'medio': metricas.get('cycle_time_medio', 0),
            'mediana': metricas.get('cycle_time_mediana', 0),
            'p75': metricas.get('cycle_time_p75', 0),
            'p95': metricas.get('cycle_time_p95', 0)
        },
        'throughput': metricas.get('throughput', 0),
        'story_points': metricas.get('story_points_medio', 0),
        'composicao': {
            'total_pessoas': metricas.get('total_pessoas', 0),
            'total_squads': metricas.get('total_squads', 0)
        }
    }
    
    # Análise Diagnóstica
    diagnosticos = []
    
    # Diagnóstico de Fluxo
    lt_medio = metricas.get('lead_time_medio', 0)
    if lt_medio > 15:  # Exemplo de threshold
        diagnosticos.append({
            'categoria': 'fluxo',
            'problema': 'Lead time elevado',
            'severidade': 'alta',
            'evidencias': [
                f"Lead time médio: {lt_medio:.1f} dias",
                f"Lead time mediano: {metricas.get('lead_time_mediana', 0):.1f} dias",
                f"Lead time P95: {metricas.get('lead_time_p95', 0):.1f} dias"
            ],
            'impacto': 'Atrasos nas entregas e aumento de custos'
        })
    
    # Diagnóstico de Eficiência
    throughput = metricas.get('throughput', 0)
    if throughput < 100:
        diagnosticos.append({
            'categoria': 'eficiencia',
            'problema': 'Throughput abaixo do esperado',
            'severidade': 'média',
            'evidencias': [
                f"Throughput atual: {throughput} entregas",
                f"Story points médios: {metricas.get('story_points_medio', 0):.1f}",
                f"Total de pessoas: {metricas.get('total_pessoas', 0)}"
            ],
            'impacto': 'Subutilização de recursos e atrasos'
        })
    
    # Análise Preditiva
    previsoes = []
    
    # Previsão de Throughput
    if throughput < 100:
        previsoes.append({
            'categoria': 'throughput',
            'cenario': 'Queda na produtividade',
            'probabilidade': 'alta',
            'evidencias': [
                f"Throughput atual: {throughput}",
                f"Story points médios: {metricas.get('story_points_medio', 0):.1f}",
                f"Total de pessoas: {metricas.get('total_pessoas', 0)}"
            ],
            'impacto_previsto': 'Redução na capacidade de entrega'
        })
    
    # Análise Prescritiva
    recomendacoes = []
    
    # Recomendações baseadas em diagnósticos
    for diag in diagnosticos:
        if diag['categoria'] == 'fluxo':
            recomendacoes.append({
                'categoria': 'fluxo',
                'acao': 'Implementar Kanban com limites de WIP',
                'prioridade': 'alta',
                'justificativa': 'Reduzir lead time e melhorar fluxo',
                'passos': [
                    'Definir limites de WIP por coluna',
                    'Implementar políticas de pull',
                    'Monitorar métricas de fluxo'
                ],
                'beneficios_esperados': [
                    'Redução do lead time',
                    'Maior previsibilidade',
                    'Melhor utilização de recursos'
                ]
            })
        elif diag['categoria'] == 'eficiencia':
            recomendacoes.append({
                'categoria': 'eficiencia',
                'acao': 'Otimizar processo de desenvolvimento',
                'prioridade': 'média',
                'justificativa': 'Aumentar eficiência do fluxo',
                'passos': [
                    'Mapear gargalos no processo',
                    'Implementar práticas de engenharia ágil',
                    'Automatizar tarefas repetitivas'
                ],
                'beneficios_esperados': [
                    'Maior eficiência',
                    'Redução de retrabalho',
                    'Melhor qualidade'
                ]
            })
    
    # Identificação de Riscos
    riscos = []
    
    # Risco de Qualidade
    story_points = metricas.get('story_points_medio', 0)
    if story_points > 5:
        riscos.append({
            'categoria': 'qualidade',
            'descricao': 'Complexidade elevada nas entregas',
            'probabilidade': 'alta',
            'impacto': 'alto',
            'evidencias': [
                f"Story points médios: {story_points:.1f}",
                f"Throughput: {throughput} entregas",
                f"Lead time médio: {lt_medio:.1f} dias"
            ],
            'mitigacao': [
                'Implementar práticas de TDD',
                'Melhorar processo de code review',
                'Aumentar cobertura de testes'
            ]
        })
    
    # Identificação de Oportunidades
    oportunidades = []
    
    # Oportunidade de Melhoria
    if throughput < 100:
        oportunidades.append({
            'categoria': 'eficiencia',
            'descricao': 'Potencial de melhoria na eficiência',
            'beneficio_esperado': 'alto',
            'evidencias': [
                f"Throughput atual: {throughput}",
                f"Story points médios: {story_points:.1f}",
                f"Total de pessoas: {metricas.get('total_pessoas', 0)}"
            ],
            'acoes_sugeridas': [
                'Implementar práticas de engenharia ágil',
                'Otimizar processo de desenvolvimento',
                'Automatizar tarefas repetitivas'
            ]
        })
    
    # Planos de Ação
    planos_acao = []
    
    # Plano baseado em recomendações
    for rec in recomendacoes:
        if rec['prioridade'] == 'alta':
            planos_acao.append({
                'titulo': rec['acao'],
                'objetivo': rec['justificativa'],
                'acoes': rec['passos'],
                'responsavel': 'Time de Agilidade',
                'prazo': '30 dias',
                'metricas_sucesso': [
                    'Redução do lead time em 20%',
                    'Aumento da eficiência em 15%',
                    'Redução da taxa de retrabalho em 25%'
                ]
            })
    
    # Atualiza a análise com todos os componentes
    analise.update({
        'diagnosticos': diagnosticos,
        'previsoes': previsoes,
        'recomendacoes': recomendacoes,
        'riscos': riscos,
        'oportunidades': oportunidades,
        'planos_acao': planos_acao
    })
    
    return analise

def formatar_analise_consultiva(analise: Dict) -> str:
    """
    Formata a análise consultiva em um relatório estruturado.
    """
    report = []
    
    # Cabeçalho
    report.append("=== Análise Consultiva ===")
    report.append(f"Data: {datetime.fromisoformat(analise['timestamp']).strftime('%d/%m/%Y %H:%M')}")
    report.append("\n")
    
    # Métricas Principais
    report.append("## Métricas Principais")
    metricas = analise['metricas_principais']
    
    # Lead Time
    if 'lead_time' in metricas:
        lt = metricas['lead_time']
        report.append("\n### Lead Time")
        report.append(f"- Médio: {lt['medio']:.1f} dias")
        report.append(f"- Mediano: {lt['mediana']:.1f} dias")
        report.append(f"- P75: {lt['p75']:.1f} dias")
        report.append(f"- P95: {lt['p95']:.1f} dias")
    
    # Cycle Time
    if 'cycle_time' in metricas:
        ct = metricas['cycle_time']
        report.append("\n### Cycle Time")
        report.append(f"- Médio: {ct['medio']:.1f} dias")
        report.append(f"- Mediano: {ct['mediana']:.1f} dias")
        report.append(f"- P75: {ct['p75']:.1f} dias")
        report.append(f"- P95: {ct['p95']:.1f} dias")
    
    # Throughput e Story Points
    report.append("\n### Produtividade")
    report.append(f"- Throughput: {metricas['throughput']} entregas")
    report.append(f"- Story Points Médios: {metricas['story_points']:.1f}")
    
    # Composição
    if 'composicao' in metricas:
        comp = metricas['composicao']
        report.append("\n### Composição")
        report.append(f"- Total de Pessoas: {comp['total_pessoas']}")
        report.append(f"- Total de Squads: {comp['total_squads']}")
        if comp['total_squads'] > 0:
            media_pessoas = comp['total_pessoas'] / comp['total_squads']
            report.append(f"- Média de Pessoas por Squad: {media_pessoas:.1f}")
    
    report.append("\n")
    
    # Diagnósticos
    if analise['diagnosticos']:
        report.append("## Diagnósticos")
        for diag in analise['diagnosticos']:
            report.append(f"\n### {diag['problema']} ({diag['severidade'].upper()})")
            report.append("Evidências:")
            for ev in diag['evidencias']:
                report.append(f"- {ev}")
            report.append(f"Impacto: {diag['impacto']}")
    report.append("\n")
    
    # Previsões
    if analise['previsoes']:
        report.append("## Previsões")
        for prev in analise['previsoes']:
            report.append(f"\n### {prev['cenario']} ({prev['probabilidade'].upper()})")
            report.append("Evidências:")
            for ev in prev['evidencias']:
                report.append(f"- {ev}")
            report.append(f"Impacto Previsto: {prev['impacto_previsto']}")
    report.append("\n")
    
    # Recomendações
    if analise['recomendacoes']:
        report.append("## Recomendações")
        for rec in analise['recomendacoes']:
            report.append(f"\n### {rec['acao']} ({rec['prioridade'].upper()})")
            report.append(f"Justificativa: {rec['justificativa']}")
            report.append("Passos:")
            for passo in rec['passos']:
                report.append(f"- {passo}")
            report.append("Benefícios Esperados:")
            for benef in rec['beneficios_esperados']:
                report.append(f"- {benef}")
    report.append("\n")
    
    # Riscos
    if analise['riscos']:
        report.append("## Riscos")
        for risco in analise['riscos']:
            report.append(f"\n### {risco['descricao']} ({risco['probabilidade'].upper()}/{risco['impacto'].upper()})")
            report.append("Evidências:")
            for ev in risco['evidencias']:
                report.append(f"- {ev}")
            report.append("Mitigação:")
            for mit in risco['mitigacao']:
                report.append(f"- {mit}")
    report.append("\n")
    
    # Oportunidades
    if analise['oportunidades']:
        report.append("## Oportunidades")
        for op in analise['oportunidades']:
            report.append(f"\n### {op['descricao']} ({op['beneficio_esperado'].upper()})")
            report.append("Evidências:")
            for ev in op['evidencias']:
                report.append(f"- {ev}")
            report.append("Ações Sugeridas:")
            for acao in op['acoes_sugeridas']:
                report.append(f"- {acao}")
    report.append("\n")
    
    # Planos de Ação
    if analise['planos_acao']:
        report.append("## Planos de Ação")
        for plano in analise['planos_acao']:
            report.append(f"\n### {plano['titulo']}")
            report.append(f"Objetivo: {plano['objetivo']}")
            report.append(f"Responsável: {plano['responsavel']}")
            report.append(f"Prazo: {plano['prazo']}")
            report.append("Ações:")
            for acao in plano['acoes']:
                report.append(f"- {acao}")
            report.append("Métricas de Sucesso:")
            for metrica in plano['metricas_sucesso']:
                report.append(f"- {metrica}")
    
    return "\n".join(report)

def printar_tribos_squads():
    import logging
    logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s', handlers=[logging.StreamHandler(), logging.FileHandler('logs/depuracao_printar_tribos.log')])
    analises = executar_pipeline()
    estrutura = None
    for a in analises:
        if isinstance(a, dict) and a.get('tipo') == 'estrutura_org':
            estrutura = a.get('dados', {})
            break
    print('Tribos:', list(estrutura['tribos'].keys()) if estrutura and 'tribos' in estrutura else [])
    print('Squads:', list(estrutura['squads'].keys()) if estrutura and 'squads' in estrutura else [])

def cruzar_dados_completo(df_maturidade, df_alocacao, df_executivo):
    """
    Cruza os dados dos três arquivos conforme o relacionamento de chaves descrito pelo usuário.
    Retorna DataFrame cruzado com métricas do Executivo associadas a tribos e squads reais.
    """
    import pandas as pd
    # Normalizar nomes de tribo e squad
    df_maturidade['tribo_norm'] = df_maturidade['Tribo'].str.strip().str.lower().fillna('')
    df_alocacao['tribe_norm'] = df_alocacao['tribe'].str.strip().str.lower().fillna('')
    df_alocacao['squad_norm'] = df_alocacao['squad'].str.strip().str.lower().fillna('')
    # Filtrar apenas pessoas ativas
    df_alocacao_ativas = df_alocacao[df_alocacao['endDate'].isna() | (pd.to_datetime(df_alocacao['endDate'], errors='coerce') > pd.Timestamp.now())].copy()
    # Merge Maturidade <-> Alocação por nome normalizado da tribo
    merged = pd.merge(df_maturidade, df_alocacao_ativas, left_on='tribo_norm', right_on='tribe_norm', how='inner', suffixes=('_mat', '_aloc'))
    # Gerar chaves compostas para cruzar com Executivo
    merged['Ano'] = merged['Ano'].astype(str)
    merged['Quarter'] = merged['Quarter'].astype(str)
    merged['Chave_DataTribo'] = merged['Ano'] + merged['Quarter'] + merged['tribeID'].astype(str)
    merged['Chave_DataSquad'] = merged['Ano'] + merged['Quarter'] + merged['squadID'].astype(str)
    # Normalizar chaves no Executivo
    df_executivo['Chave_DataTribo'] = df_executivo['PBI_Concuidos_Executivo[Chave_DataTribo]'].astype(str).str.strip()
    df_executivo['Chave_DataSquad'] = df_executivo['PBI_Concuidos_Executivo[Chave_DataSquad]'].astype(str).str.strip()
    # Merge com Executivo por Chave_DataTribo e Chave_DataSquad
    merged = pd.merge(merged, df_executivo, left_on='Chave_DataTribo', right_on='Chave_DataTribo', how='left', suffixes=('', '_exec'))
    merged = pd.merge(merged, df_executivo, left_on='Chave_DataSquad', right_on='Chave_DataSquad', how='left', suffixes=('', '_exec_squad'))
    return merged

def gerar_estrutura_e_insights(df_cruzado):
    """
    Popula a estrutura organizacional, calcula métricas e gera insights para cada tribo e squad a partir do DataFrame cruzado.
    Retorna estrutura, métricas e insights por tribo e squad.
    """
    import numpy as np
    estrutura = {'tribos': {}, 'squads': {}, 'pessoas': set()}
    insights_tribos = {}
    insights_squads = {}
    # Tribos
    for tribo in df_cruzado['Tribo'].dropna().unique():
        df_tribo = df_cruzado[df_cruzado['Tribo'] == tribo]
        squads = df_tribo['squad'].dropna().unique().tolist()
        pessoas = df_tribo['person'].dropna().unique().tolist()
        estrutura['tribos'][tribo] = {'squads': squads, 'pessoas': pessoas}
        estrutura['pessoas'].update(pessoas)
        # Métricas de fluxo
        lead_times = df_tribo['[SumLead_Time]'].dropna().astype(float).tolist() if '[SumLead_Time]' in df_tribo else []
        cycle_times = df_tribo['[SumCycle_Time]'].dropna().astype(float).tolist() if '[SumCycle_Time]' in df_tribo else []
        story_points = df_tribo['[SumStory_Points]'].dropna().astype(float).tolist() if '[SumStory_Points]' in df_tribo else []
        throughput = df_tribo['PBI_Concuidos_Executivo[Key]'].nunique() if 'PBI_Concuidos_Executivo[Key]' in df_tribo else 0
        insights_tribos[tribo] = {
            'lead_time_medio': float(np.mean(lead_times)) if lead_times else 0,
            'lead_time_mediana': float(np.median(lead_times)) if lead_times else 0,
            'lead_time_p75': float(np.percentile(lead_times, 75)) if lead_times else 0,
            'lead_time_p95': float(np.percentile(lead_times, 95)) if lead_times else 0,
            'cycle_time_medio': float(np.mean(cycle_times)) if cycle_times else 0,
            'cycle_time_mediana': float(np.median(cycle_times)) if cycle_times else 0,
            'cycle_time_p75': float(np.percentile(cycle_times, 75)) if cycle_times else 0,
            'cycle_time_p95': float(np.percentile(cycle_times, 95)) if cycle_times else 0,
            'story_points_medio': float(np.mean(story_points)) if story_points else 0,
            'throughput': throughput,
            'total_pessoas': len(pessoas),
            'total_squads': len(squads)
        }
    # Squads
    for squad in df_cruzado['squad'].dropna().unique():
        df_squad = df_cruzado[df_cruzado['squad'] == squad]
        pessoas = df_squad['person'].dropna().unique().tolist()
        tribos = df_squad['Tribo'].dropna().unique().tolist()
        estrutura['squads'][squad] = {'tribos': tribos, 'pessoas': pessoas}
        estrutura['pessoas'].update(pessoas)
        lead_times = df_squad['[SumLead_Time]'].dropna().astype(float).tolist() if '[SumLead_Time]' in df_squad else []
        cycle_times = df_squad['[SumCycle_Time]'].dropna().astype(float).tolist() if '[SumCycle_Time]' in df_squad else []
        story_points = df_squad['[SumStory_Points]'].dropna().astype(float).tolist() if '[SumStory_Points]' in df_squad else []
        throughput = df_squad['PBI_Concuidos_Executivo[Key]'].nunique() if 'PBI_Concuidos_Executivo[Key]' in df_squad else 0
        insights_squads[squad] = {
            'lead_time_medio': float(np.mean(lead_times)) if lead_times else 0,
            'lead_time_mediana': float(np.median(lead_times)) if lead_times else 0,
            'lead_time_p75': float(np.percentile(lead_times, 75)) if lead_times else 0,
            'lead_time_p95': float(np.percentile(lead_times, 95)) if lead_times else 0,
            'cycle_time_medio': float(np.mean(cycle_times)) if cycle_times else 0,
            'cycle_time_mediana': float(np.median(cycle_times)) if cycle_times else 0,
            'cycle_time_p75': float(np.percentile(cycle_times, 75)) if cycle_times else 0,
            'cycle_time_p95': float(np.percentile(cycle_times, 95)) if cycle_times else 0,
            'story_points_medio': float(np.mean(story_points)) if story_points else 0,
            'throughput': throughput,
            'total_pessoas': len(pessoas),
            'total_tribos': len(tribos)
        }
    estrutura['pessoas'] = list(estrutura['pessoas'])
    return estrutura, insights_tribos, insights_squads

# No pipeline, após carregar os dados:
# dados = carregar_dados()
# df_cruzado = cruzar_dados_completo(dados['maturidade'], dados['alocacao'], dados['executivo'])
# (usar df_cruzado para enriquecer estrutura e análises)

# No pipeline, após cruzar os dados:
# estrutura, insights_tribos, insights_squads = gerar_estrutura_e_insights(df_cruzado)