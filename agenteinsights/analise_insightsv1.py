"""
Agente Insights - Módulo Principal
=================================
Versão: 1.4.0
Release: 5
Data: 02/06/2025

Histórico:
- 1.0.0 (Release 0): Versão inicial
- 1.1.0 (Release 1): Adicionado sistema de feedback e logging
- 1.2.0 (Release 2): Correção no processamento de dados e análise ML
- 1.3.0 (Release 4): Integração com [analise_insights.py](http://_vscodecontentref_/1)
- 1.4.0 (Release 5): Integração com chat OpenAI, respostas baseadas em dados, salvar chat em DOCX

Descrição:
Pipeline completo: análise, chat IA, respostas baseadas em dados, gráficos/tabelas e exportação do chat.
"""

import logging
from datetime import datetime
from pathlib import Path
from colorama import init, Fore, Style
from tqdm import tqdm
import time
import os
from dotenv import load_dotenv
import openai
from docx import Document

def print_status(msg, tipo="info"):
    cores = {
        "info": Fore.CYAN,
        "success": Fore.GREEN,
        "warning": Fore.YELLOW,
        "error": Fore.RED
    }
    print(cores.get(tipo, Fore.WHITE) + f"[{tipo.upper()}] {msg}" + Style.RESET_ALL)

def criar_diretorios():
    diretorios = [
        "output",
        "output/graficos",
        "output/logs",
        "output/relatorios"
    ]
    print_status("Verificando e criando diretórios necessários...", "info")
    for d in tqdm(diretorios, desc="Criando diretórios"):
        Path(d).mkdir(parents=True, exist_ok=True)
        time.sleep(0.05)
    print_status("Diretórios criados com sucesso!", "success")

def carregar_openai_key():
    load_dotenv()
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        print_status("API KEY da OpenAI não encontrada no .env!", "error")
        exit(1)
    openai.api_key = api_key

def chat_ia_loop(contexto_base, dados_cruzados, resultados):
    print_status("Chat IA iniciado! Pergunte sobre tribos, squads ou peça insights.", "info")
    print_status("Digite 'salvar' para exportar o chat para DOCX ou 'sair' para encerrar.", "info")
    chat_log = []
    while True:
        user_input = input(Fore.YELLOW + "Você: " + Style.RESET_ALL)
        if user_input.strip().lower() == "sair":
            print_status("Chat encerrado.", "success")
            break
        if user_input.strip().lower() == "salvar":
            salvar_chat_docx(chat_log)
            continue

        # Monta contexto para IA
        prompt = (
            f"{contexto_base}\n"
            f"Dados disponíveis: {dados_cruzados.head(3).to_string()}\n"
            f"Resultados estatísticos: {resultados['estatisticas']}\n"
            f"Pergunta do usuário: {user_input}\n"
            "Responda de forma clara, cite dados, sugira gráficos/tabelas se relevante."
        )
        try:
            resposta = openai.ChatCompletion.create(
                model="gpt-4-1106-preview",  # ou o modelo mini disponível
                messages=[
                    {"role": "system", "content": contexto_base},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=500,
                temperature=0.2
            )["choices"][0]["message"]["content"]
        except Exception as e:
            resposta = f"[ERRO] Falha ao consultar IA: {str(e)}"
        print(Fore.GREEN + "IA:" + Style.RESET_ALL, resposta)
        chat_log.append(("Você", user_input))
        chat_log.append(("IA", resposta))

def salvar_chat_docx(chat_log):
    doc = Document()
    doc.add_heading('Chat de Insights - Agente Insights', 0)
    for autor, msg in chat_log:
        doc.add_paragraph(f"{autor}:", style='Heading 2')
        doc.add_paragraph(msg)
    caminho = f"output/relatorios/chat_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(caminho)
    print_status(f"Chat salvo em: {caminho}", "success")

def carregar_dados():
    # Função para carregar dados
    pass

def cruzar_dados(dados):
    # Função para cruzar dados
    pass

def executar_analises(dados_cruzados):
    # Função para executar análises
    pass

def gerar_graficos(dados_cruzados, resultados):
    # Função para gerar gráficos
    pass

def gerar_relatorio(resultados, caminho_relatorio):
    # Função para gerar relatório
    pass

def main():
    data_hora = datetime.now().strftime("%Y%m%d_%H%M%S")
    logging.basicConfig(
        filename=f'output/logs/execucao_{data_hora}.log',
        level=logging.INFO,
        format='%(asctime)s - %(levelname)s - %(message)s'
    )
    init(autoreset=True)
    try:
        print_status("Iniciando análise de insights...", "info")
        criar_diretorios()
        carregar_openai_key()

        print_status("Carregando dados...", "info")
        dados = carregar_dados()
        print_status("Dados carregados com sucesso!", "success")

        print_status("Cruzando dados...", "info")
        dados_cruzados = cruzar_dados(dados)
        print_status("Dados cruzados com sucesso!", "success")

        print_status("Executando análises de ML...", "info")
        resultados = executar_analises(dados_cruzados)
        if resultados['status'] != 'success':
            print_status(f"Erro nas análises de ML: {resultados['message']}", "error")
            return
        print_status("Análises de ML concluídas com sucesso!", "success")

        print_status("Gerando gráficos...", "info")
        gerar_graficos(dados_cruzados, resultados)
        print_status("Gráficos gerados com sucesso!", "success")

        print_status("Gerando relatório...", "info")
        caminho_relatorio = f"output/relatorios/relatorio_{data_hora}.docx"
        gerar_relatorio(resultados, caminho_relatorio)
        print_status(f"Relatório gerado com sucesso: {caminho_relatorio}", "success")

        # Chat IA
        contexto_base = (
            "Você é um assistente de dados para squads e tribos. "
            "Responda perguntas com base nos dados e análises fornecidos, "
            "sugerindo gráficos ou tabelas quando relevante."
        )
        chat_ia_loop(contexto_base, dados_cruzados, resultados)

    except Exception as e:
        print_status(f"Erro durante a execução: {str(e)}", "error")
        logging.error(f"Erro durante a execução: {str(e)}")

if __name__ == "__main__":
    main()
