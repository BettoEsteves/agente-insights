"""
Agente Insights - Análises e Chat IA
====================================
Versão: 1.5.0
Release: 5
Data: 02/06/2025

Descrição:
Executa análises (descritiva, preditiva, diagnóstica) e fornece função para chat IA,
onde o usuário pode consultar insights sobre tribos/squads, com respostas baseadas nos dados,
gráficos e tabelas. Permite salvar o histórico do chat em DOCX.
"""

import pandas as pd
import numpy as np
import os
import logging
from sklearn.linear_model import LinearRegression
from sklearn.cluster import KMeans
from sklearn.preprocessing import StandardScaler
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from datetime import datetime
from typing import Dict, Any, List
from dotenv import load_dotenv
import openai
from openai import OpenAI  # Adicione esta linha junto com os outros imports
import unicodedata

# Configurar logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

def normalizar_coluna(col):
    # Remove acentos, espaços e deixa minúsculo
    col = unicodedata.normalize('NFKD', str(col)).encode('ASCII', 'ignore').decode('ASCII')
    return col.strip().lower().replace(' ', '')

def mapear_colunas(df, nomes_esperados):
    # Cria um dicionário de {nome_normalizado: nome_original}
    col_map = {normalizar_coluna(c): c for c in df.columns}
    resultado = {}
    for nome in nomes_esperados:
        norm = normalizar_coluna(nome)
        if norm in col_map:
            resultado[nome] = col_map[norm]
        else:
            resultado[nome] = None
    return resultado

def carregar_dados() -> Dict[str, pd.DataFrame]:
    maturidade = pd.read_excel('MaturidadeT.xlsx', sheet_name='FT_Pesquisa_Nota_Maturidade_Por', engine='openpyxl')
    alocacao = pd.read_excel('Alocacao.xlsx', sheet_name='Allocation', engine='openpyxl')
    executivo = pd.read_excel('Executivo.xlsx', sheet_name='NewBusinessAgility', engine='openpyxl')
    return {'maturidade': maturidade, 'alocacao': alocacao, 'executivo': executivo}

def cruzar_dados(dados: Dict[str, pd.DataFrame]) -> pd.DataFrame:
    # Mapeia os nomes das colunas reais para os nomes esperados
    mat_map = mapear_colunas(dados['maturidade'], ['Tribo'])
    alo_map = mapear_colunas(dados['alocacao'], ['tribe', 'squad', 'squadID', 'tribeID'])
    exe_map = mapear_colunas(dados['executivo'], [
        'PBI_Concuidos_Executivo[ID_Squad]', 'PBI_Concuidos_Executivo[ID_Tribo]'
    ])

    # Merge 1: Tribo (maturidade) <-> tribe (alocacao)
    df1 = dados['maturidade'].merge(
        dados['alocacao'],
        left_on=mat_map['Tribo'],
        right_on=alo_map['tribe'],
        how='left',
        suffixes=('', '_aloc')
    )
    # Merge 2: squadID (alocacao) <-> PBI_Concuidos_Executivo[ID_Squad] (executivo)
    df2 = df1.merge(
        dados['executivo'],
        left_on=alo_map['squadID'],
        right_on=exe_map['PBI_Concuidos_Executivo[ID_Squad]'],
        how='left',
        suffixes=('', '_exec')
    )
    # (Opcional) Merge 3: tribeID (alocacao) <-> PBI_Concuidos_Executivo[ID_Tribo] (executivo)
    # df2 = df2.merge(
    #     dados['executivo'],
    #     left_on=alo_map['tribeID'],
    #     right_on=exe_map['PBI_Concuidos_Executivo[ID_Tribo]'],
    #     how='left',
    #     suffixes=('', '_exec_tribo')
    # )
    return df2

def executar_analises(df: pd.DataFrame) -> Dict[str, Any]:
    resultados = {}
    # Análise descritiva
    resultados['estatisticas'] = df.describe(include='all').to_dict()
    # Análise preditiva (regressão)
    num_cols = df.select_dtypes(include=[np.number]).columns
    if len(num_cols) > 1:
        X = df[num_cols[1:]].fillna(0)
        y = df[num_cols[0]].fillna(0)
        reg = LinearRegression().fit(X, y)
        resultados['regressao'] = {
            'coef': reg.coef_.tolist(),
            'intercept': float(reg.intercept_),
            'r2': reg.score(X, y)
        }
    else:
        resultados['regressao'] = None
    # Análise diagnóstica (clustering)
    if len(num_cols) > 1:
        X = df[num_cols].fillna(0)
        kmeans = KMeans(n_clusters=3, random_state=42).fit(X)
        resultados['clustering'] = {
            'labels': kmeans.labels_.tolist(),
            'centroids': kmeans.cluster_centers_.tolist(),
            'inertia': float(kmeans.inertia_)
        }
    else:
        resultados['clustering'] = None
    resultados['status'] = 'success'
    return resultados

def gerar_graficos(df: pd.DataFrame, resultados: Dict[str, Any]):
    # Exemplo: histograma da primeira coluna numérica
    num_cols = df.select_dtypes(include=[np.number]).columns
    if len(num_cols) > 0:
        plt.figure(figsize=(8, 4))
        sns.histplot(df[num_cols[0]].dropna())
        plt.title(f'Histograma de {num_cols[0]}')
        plt.savefig('output/graficos/histograma.png')
        plt.close()

def chat_ia_loop(analises: Dict[str, Any]):
    # Carrega variáveis de ambiente
    load_dotenv()
    
    # Inicializa cliente OpenAI
    client = openai.OpenAI()  # Usando referência completa
    
    # Contexto base como uma mensagem de desenvolvedor
    contexto_base = {
        "role": "developer",
        "content": """
        # Identidade e Especialidades
        Você é um especialista em:
        * Gestão empresarial e organizacional
        * Agilidade e Business Agility
        * Gestão de projetos e produtos
        * Planejamento estratégico e indicadores
        
        # Instruções
        * Analise dados sobre maturidade, alocação e qualidade do trabalho
        * Forneça insights baseados em estatísticas e tendências
        * Use tom profissional para gestores
        * Sugira planos de ação e monitoramento
        * Cite dados específicos das análises fornecidas
        """
    }
    
    chat_log = []
    print("Chat IA iniciado! Pergunte sobre tribos, squads ou peça insights.")
    print("Digite 'salvar' para exportar o chat para DOCX ou 'sair' para encerrar.")
    
    while True:
        user_input = input("Você: ")
        if user_input.strip().lower() == "sair":
            print("Chat encerrado.")
            break
        if user_input.strip().lower() == "salvar":
            salvar_chat_docx(chat_log)
            continue

        try:
            # Nova sintaxe para responses.create
            response = client.responses.create(
                model="gpt-4.1",
                input=[
                    contexto_base,
                    {
                        "role": "user",
                        "content": f"""
                        # Contexto
                        Estatísticas: {analises['descritiva']}
                        Regressão: {analises['preditiva']}
                        Clustering: {analises['diagnostica']}
                        
                        # Pergunta
                        {user_input}
                        """
                    }
                ]
            )
            
            resposta = response.output_text
            
        except Exception as e:
            logging.error(f"Erro ao consultar OpenAI: {str(e)}")
            resposta = f"[ERRO] Falha ao consultar IA: {str(e)}"
        
        print("IA:", resposta)
        chat_log.append(("Você", user_input))
        chat_log.append(("IA", resposta))

def salvar_chat_docx(chat_log: List[tuple]):
    doc = Document()
    doc.add_heading('Chat de Insights - Agente Insights', 0)
    for autor, msg in chat_log:
        doc.add_paragraph(f"{autor}:", style='Heading 2')
        doc.add_paragraph(msg)
    caminho = f"output/relatorios/chat_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(caminho)
    print(f"Chat salvo em: {caminho}")

def executar_pipeline():
    dados = carregar_dados()
    dados_cruzados = cruzar_dados(dados)
    resultados = executar_analises(dados_cruzados)
    gerar_graficos(dados_cruzados, resultados)
    analises = {
        "dados": dados,
        "dados_cruzados": dados_cruzados,
        "descritiva": resultados.get("estatisticas"),
        "preditiva": resultados.get("regressao"),
        "diagnostica": resultados.get("clustering"),
        "status": resultados.get("status"),
        "mensagem": resultados.get("message", "")
    }
    return analises

# Remova todo o código abaixo desta linha
# Não execute nada automaticamente!
# O insights.py deve importar executar_pipeline e chat_ia_loop para orquestrar o fluxo.
