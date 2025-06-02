
from docx import Document
from docx.shared import Inches
import pandas as pd
import os
import logging
from datetime import datetime
from typing import Dict, Any

def gerar_docx(resultados: Dict[str, Any]) -> None:
    """
    Gera um relatório em formato DOCX com os resultados das análises.
    
    Args:
        resultados: Dicionário contendo os resultados das análises
            - descritiva: DataFrame com estatísticas descritivas
            - regressao: Dict com coeficientes e R² da regressão
            - clusters: Dict com centros e inertia dos clusters
            - correlacao: DataFrame com matriz de correlação
    """
    try:
        # Criar diretório para relatórios se não existir
        os.makedirs('output/relatorios', exist_ok=True)
        
        # Iniciar documento
        logging.info("Iniciando geração do relatório...")
        doc = Document()
        
        # Cabeçalho
        doc.add_heading('Relatório Final de Análises', 0)
        doc.add_paragraph(f'Gerado em: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}')
        
        # Análise Descritiva
        logging.info("Adicionando análise descritiva...")
        doc.add_heading('Análise Descritiva', level=1)
        doc.add_paragraph(resultados['descritiva'].to_string())
        
        # Análise Preditiva
        logging.info("Adicionando análise preditiva...")
        doc.add_heading('Análise Preditiva (Regressão Linear)', level=1)
        regressao = resultados['regressao']
        doc.add_paragraph('Coeficientes do modelo:')
        for var, coef in regressao['coeficientes'].items():
            doc.add_paragraph(f'- {var}: {coef:.4f}', style='List Bullet')
        doc.add_paragraph(f'R²: {regressao["r2"]:.4f}')
        
        # Análise Prescritiva
        logging.info("Adicionando análise prescritiva...")
        doc.add_heading('Análise Prescritiva (Clustering)', level=1)
        clusters = resultados['clusters']
        doc.add_paragraph(f'Inércia total dos clusters: {clusters["inertia"]:.2f}')
        doc.add_paragraph('Centros dos clusters:')
        for i, centro in enumerate(clusters['centros']):
            doc.add_paragraph(f'Cluster {i}: {centro}', style='List Bullet')
        
        # Análise Diagnóstica
        logging.info("Adicionando análise diagnóstica...")
        doc.add_heading('Análise Diagnóstica (Correlação)', level=1)
        doc.add_paragraph(resultados['correlacao'].round(4).to_string())
        
        # Gráficos
        logging.info("Adicionando gráficos...")
        doc.add_heading('Visualizações', level=1)
        
        # Heatmap de Correlação
        doc.add_heading('Heatmap de Correlação', level=2)
        doc.add_picture('output/graficos/correlacao_heatmap.png', 
                       width=Inches(6))
        
        # Gráfico de Clusters
        doc.add_heading('Análise de Clusters', level=2)
        doc.add_picture('output/graficos/clusters_scatterplot.png', 
                       width=Inches(6))
        
        # Salvar documento
        nome_arquivo = f'output/relatorios/relatorio_final_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
        doc.save(nome_arquivo)
        logging.info(f"Relatório salvo com sucesso em: {nome_arquivo}")
        
    except Exception as e:
        logging.error(f"Erro ao gerar relatório: {str(e)}")
        raise
